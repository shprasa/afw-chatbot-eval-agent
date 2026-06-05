"""Deliverables for Claude v1 vs Claude v10 comparison (120 personas).

Outputs:
  - powerbi_export/AFW_Claude_v1_vs_v10_Data.xlsx
  - powerbi_export/AFW_Claude_v1_vs_v10_Report.docx
"""
from __future__ import annotations

import json
import os
import tempfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from sklearn.metrics import classification_report, confusion_matrix
from statsmodels.stats.contingency_tables import mcnemar

DESK = Path(__file__).resolve().parent
ARTIFACTS = DESK / "artifacts"
REPORTS = DESK / "reports"
SRC_120 = DESK / "AFW_120_User_Test_Cases_Original_Style_Short.xlsx"
OUT_DIR = DESK / "powerbi_export"
LOCAL = Path(os.environ.get("LOCALAPPDATA", "")) / "AFW_powerbi_export"

V1_PRED = ARTIFACTS / "chatbot_live_predictions_original_style_short_claude_v1.csv"
V10_PRED = ARTIFACTS / "chatbot_live_predictions_original_style_short_claude.csv"
V1_ACC = REPORTS / "chatbot_live_accuracy_original_style_short_claude_v1.json"
V10_ACC = REPORTS / "chatbot_live_accuracy_original_style_short_claude.json"
V1_FAIL_MD = REPORTS / "chatbot_live_failure_analysis_original_style_short_claude_v1.md"
V1_PROMPT_MD = REPORTS / "chatbot_live_prompt_improvements_original_style_short_claude_v1.md"
V10_FAIL_MD = REPORTS / "chatbot_live_failure_analysis_original_style_short_claude.md"
V10_PROMPT_MD = REPORTS / "chatbot_live_prompt_improvements_original_style_short_claude.md"

XLSX_OUT = "AFW_Claude_v1_vs_v10_Data.xlsx"
DOCX_OUT = "AFW_Claude_v1_vs_v10_Report.docx"

QN = list(range(1, 9))
ROMAN = ["i", "ii", "iii", "iv", "v", "vi", "vii", "viii"]
CODE_TO_TEXT = {
    1: "eligible",
    0: "ineligible",
    2: "insufficient_information",
    3: "manual_review",
    "1": "eligible",
    "0": "ineligible",
    "2": "insufficient_information",
    "3": "manual_review",
}
CLASSES = ["eligible", "ineligible", "manual_review", "insufficient_information"]


def code_to_text(v) -> str:
    if pd.isna(v):
        return ""
    if isinstance(v, str) and not v.replace(".", "").isdigit():
        return v.strip()
    try:
        return CODE_TO_TEXT.get(int(float(v)), str(v))
    except (TypeError, ValueError):
        return str(v)


def load_paired() -> tuple[pd.DataFrame, pd.DataFrame]:
    v1 = pd.read_csv(V1_PRED).drop_duplicates("persona_id")
    v10 = pd.read_csv(V10_PRED).drop_duplicates("persona_id")
    common = sorted(set(v1["persona_id"]) & set(v10["persona_id"]))
    v1 = v1[v1["persona_id"].isin(common)].sort_values("persona_id")
    v10 = v10[v10["persona_id"].isin(common)].sort_values("persona_id")
    return v1, v10


def build_120_sheet(pred_csv: Path, persona_ids: list[str]) -> pd.DataFrame:
    src = pd.read_excel(SRC_120, sheet_name="120_test_cases", engine="openpyxl")
    src.columns = src.columns.str.lower()
    pred = pd.read_csv(pred_csv).drop_duplicates(subset="persona_id", keep="last")
    pred = pred[pred["persona_id"].isin(persona_ids)]

    rows: list[dict] = []
    for _, s in src.iterrows():
        pid = s["persona_id"]
        if pid not in persona_ids:
            continue
        p = pred[pred["persona_id"] == pid]
        p = p.iloc[0] if len(p) else None
        row: dict = {
            "persona_id": pid,
            "engineered_for (truth final)": s.get("engineered_for", ""),
        }
        for n, r in zip(QN, ROMAN):
            row[f"q{n} user input message"] = s.get(f"{r}. user input message", "")
            row[f"q{n} manual label"] = s.get(f"{r}. manual label", "")
            row[f"q{n} predicted"] = (
                code_to_text(p.get(f"pred_q{n}")) if p is not None else ""
            )
        truth_final = str(s.get("engineered_for", "")).strip()
        pred_final = code_to_text(p.get("predicted_label")) if p is not None else ""
        row["final manual label"] = truth_final
        row["final predicted"] = pred_final
        row["final match"] = (
            "" if not pred_final else ("TRUE" if truth_final == pred_final else "FALSE")
        )
        rows.append(row)
    return pd.DataFrame(rows)


def mcnemar_stats(v1_df: pd.DataFrame, v10_df: pd.DataFrame) -> dict:
    m = v1_df[["persona_id", "label_match"]].merge(
        v10_df[["persona_id", "label_match"]], on="persona_id", suffixes=("_v1", "_v10")
    )
    m["v1_ok"] = m["label_match_v1"].astype(bool)
    m["v10_ok"] = m["label_match_v10"].astype(bool)
    v1_only = int((m["v1_ok"] & ~m["v10_ok"]).sum())
    v10_only = int((~m["v1_ok"] & m["v10_ok"]).sum())
    both = int((m["v1_ok"] & m["v10_ok"]).sum())
    neither = int((~m["v1_ok"] & ~m["v10_ok"]).sum())
    table = [[both, v10_only], [v1_only, neither]]
    res = mcnemar(table, exact=True)
    return {
        "n_paired": len(m),
        "v1_accuracy": float(m["v1_ok"].mean()),
        "v10_accuracy": float(m["v10_ok"].mean()),
        "v1_only_correct": v1_only,
        "v10_only_correct": v10_only,
        "both_correct": both,
        "both_wrong": neither,
        "discordant": v1_only + v10_only,
        "mcnemar_p_value": float(res.pvalue),
        "significant_at_0_05": bool(res.pvalue < 0.05),
    }


def confusion_metrics(df: pd.DataFrame, model_label: str) -> tuple[pd.DataFrame, pd.DataFrame]:
    y_true = df["truth_label"]
    y_pred = df["predicted_label"]
    cm = confusion_matrix(y_true, y_pred, labels=CLASSES)
    cm_df = pd.DataFrame(
        cm,
        index=[f"truth_{c}" for c in CLASSES],
        columns=[f"pred_{c}" for c in CLASSES],
    )
    cm_df.index.name = model_label
    rep = classification_report(
        y_true, y_pred, labels=CLASSES, output_dict=True, zero_division=0
    )
    rows = []
    for c in CLASSES:
        r = rep[c]
        rows.append(
            {
                "model": model_label,
                "class": c,
                "precision": r["precision"],
                "recall": r["recall"],
                "f1": r["f1-score"],
                "support": int(r["support"]),
            }
        )
    macro = rep["macro avg"]
    rows.append(
        {
            "model": model_label,
            "class": "macro_avg",
            "precision": macro["precision"],
            "recall": macro["recall"],
            "f1": macro["f1-score"],
            "support": int(macro["support"]),
        }
    )
    return cm_df, pd.DataFrame(rows)


def pct(v: float) -> str:
    return f"{100.0 * v:.1f}%"


def md_to_docx_sections(doc: Document, md_text: str) -> None:
    for line in md_text.splitlines():
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("---"):
            continue
        else:
            doc.add_paragraph(line.strip())


def atomic_save_docx(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fd, tmp_name = tempfile.mkstemp(suffix=".docx", dir=str(path.parent))
    os.close(fd)
    tmp = Path(tmp_name)
    try:
        doc.save(str(tmp))
        if path.exists():
            path.unlink()
        tmp.replace(path)
    finally:
        if tmp.exists() and not path.exists():
            tmp.unlink(missing_ok=True)


def build_docx(stats: dict, cm_v1: pd.DataFrame, cm_v10: pd.DataFrame, pr_df: pd.DataFrame) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading("AFW Chatbot Evaluation: Claude v1 vs Claude v10", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Dataset: AFW_120_User_Test_Cases_Original_Style_Short (F001-F120) | "
        "Backend: angel-flight-chatbot-claude | "
        "Prompt A/B: System Prompt v1 vs v10"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Executive summary", level=1)
    diff_pp = (stats["v10_accuracy"] - stats["v1_accuracy"]) * 100
    fav = "v1" if stats["v1_accuracy"] > stats["v10_accuracy"] else "v10"
    sig = (
        "statistically significant"
        if stats["significant_at_0_05"]
        else "not statistically significant"
    )
    doc.add_paragraph(
        f"On paired final-outcome labels (n={stats['n_paired']}), Claude v1 achieved "
        f"{pct(stats['v1_accuracy'])} ({int(stats['v1_accuracy'] * stats['n_paired'])}/120) "
        f"and Claude v10 achieved {pct(stats['v10_accuracy'])} "
        f"({int(stats['v10_accuracy'] * stats['n_paired'])}/120). "
        f"Difference (v10 - v1): {diff_pp:+.1f} pp (favors {fav}). "
        f"McNemar exact test: p={stats['mcnemar_p_value']:.4f} — {sig} at alpha=0.05."
    )
    doc.add_paragraph(
        "v1 slightly outperforms v10 on this cohort, mainly on manual_review recall "
        "(v1 holds borderline cases as MR more often; v10 collapses more MR truth into "
        "insufficient_information)."
    )

    doc.add_heading("2. McNemar test (Claude v1 vs Claude v10)", level=1)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"
    for k, v in [
        ("Paired personas", str(stats["n_paired"])),
        ("Claude v1 accuracy", pct(stats["v1_accuracy"])),
        ("Claude v10 accuracy", pct(stats["v10_accuracy"])),
        ("Difference (v10 - v1)", f"{diff_pp:+.1f} pp"),
        ("Both correct", str(stats["both_correct"])),
        ("Claude v1 only correct", str(stats["v1_only_correct"])),
        ("Claude v10 only correct", str(stats["v10_only_correct"])),
        ("Both wrong", str(stats["both_wrong"])),
        ("Discordant pairs", str(stats["discordant"])),
        ("McNemar exact two-sided p-value", f"{stats['mcnemar_p_value']:.4f}"),
        ("Significant at alpha=0.05", "Yes" if stats["significant_at_0_05"] else "No"),
    ]:
        cells = tbl.add_row().cells
        cells[0].text = k
        cells[1].text = v

    doc.add_paragraph(
        "Contingency (rows=Claude v1 correct, cols=Claude v10 correct): "
        f"[[{stats['both_correct']}, {stats['v10_only_correct']}], "
        f"[{stats['v1_only_correct']}, {stats['both_wrong']}]]"
    )

    doc.add_heading("3. Confusion matrices and per-class metrics", level=1)
    for name, cm in [("Claude v1", cm_v1), ("Claude v10", cm_v10)]:
        doc.add_heading(name, level=2)
        doc.add_paragraph("Confusion matrix (rows=truth, cols=pred):")
        t = doc.add_table(rows=1, cols=len(CLASSES) + 1)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        hdr[0].text = "truth \\ pred"
        for j, c in enumerate(CLASSES):
            hdr[j + 1].text = c
        for i, truth in enumerate(CLASSES):
            cells = t.add_row().cells
            cells[0].text = truth
            for j in range(len(CLASSES)):
                cells[j + 1].text = str(int(cm.iloc[i, j]))
        sub_pr = pr_df[pr_df["model"] == name]
        doc.add_paragraph("Precision / recall / F1:")
        t2 = doc.add_table(rows=1, cols=4)
        t2.style = "Light Grid Accent 1"
        for j, htxt in enumerate(["Class", "Precision", "Recall", "F1"]):
            t2.rows[0].cells[j].text = htxt
        for _, r in sub_pr.iterrows():
            if r["class"] == "macro_avg":
                continue
            cells = t2.add_row().cells
            cells[0].text = str(r["class"])
            cells[1].text = f"{r['precision']:.3f}"
            cells[2].text = f"{r['recall']:.3f}"
            cells[3].text = f"{r['f1']:.3f}"
        macro = sub_pr[sub_pr["class"] == "macro_avg"].iloc[0]
        doc.add_paragraph(
            f"Macro avg: P={macro['precision']:.3f}  R={macro['recall']:.3f}  "
            f"F1={macro['f1']:.3f}"
        )

    doc.add_heading("4. Headline accuracy by prompt version", level=1)
    for label, acc_path in [("Claude v1", V1_ACC), ("Claude v10", V10_ACC)]:
        if not acc_path.is_file():
            continue
        acc = json.loads(acc_path.read_text(encoding="utf-8"))
        doc.add_heading(label, level=2)
        doc.add_paragraph(
            f"Final outcome accuracy (report file): {pct(acc['final_outcome_accuracy'])} "
            f"({acc.get('correct', 'n/a')}/{acc.get('n_scored', 'n/a')})."
        )

    doc.add_heading("5. Failure analysis — Claude v1 (rationale-driven)", level=1)
    if V1_FAIL_MD.is_file():
        md_to_docx_sections(doc, V1_FAIL_MD.read_text(encoding="utf-8"))

    doc.add_heading("6. Failure analysis — Claude v10 (rationale-driven)", level=1)
    if V10_FAIL_MD.is_file():
        md_to_docx_sections(doc, V10_FAIL_MD.read_text(encoding="utf-8"))

    doc.add_heading("7. System prompt recommendations — Claude v1", level=1)
    doc.add_paragraph("Source: artifacts/prompt_extract_v1.txt")
    if V1_PROMPT_MD.is_file():
        md_to_docx_sections(doc, V1_PROMPT_MD.read_text(encoding="utf-8"))

    doc.add_heading("8. System prompt recommendations — Claude v10", level=1)
    doc.add_paragraph("Source: artifacts/prompt_extract_v10.txt (or deployed v10 prompt)")
    if V10_PROMPT_MD.is_file():
        md_to_docx_sections(doc, V10_PROMPT_MD.read_text(encoding="utf-8"))

    doc.add_heading("9. v1 vs v10 prompt delta (actionable)", level=1)
    for line in [
        "v10 should retain v1 strength on insufficient_information (do not over-escalate to manual_review).",
        "v10 should recover v1 manual_review recall: borderline mobility, stability, and companion cases need MR not insufficient_information.",
        "Both versions: gate eligible on full checkpoint satisfaction; use ineligible when disqualifier is confirmed.",
        "Require per-turn *_rationale and outcome_notes before each eligibility_outcome update.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("10. Data companion", level=1)
    doc.add_paragraph(
        f"Workbook: {XLSX_OUT} — sheets Claude_v1_120, Claude_v10_120, McNemar_Summary, "
        "Confusion_Metrics, Confusion_Claude_v1, Confusion_Claude_v10."
    )

    for out_dir in (OUT_DIR, LOCAL):
        if out_dir:
            atomic_save_docx(doc, out_dir / DOCX_OUT)


def main() -> None:
    v1_df, v10_df = load_paired()
    persona_ids = sorted(v1_df["persona_id"].tolist())
    stats = mcnemar_stats(v1_df, v10_df)
    cm_v1, pr_v1 = confusion_metrics(v1_df, "Claude v1")
    cm_v10, pr_v10 = confusion_metrics(v10_df, "Claude v10")
    pr_all = pd.concat([pr_v1, pr_v10], ignore_index=True)

    v1_sheet = build_120_sheet(V1_PRED, persona_ids)
    v10_sheet = build_120_sheet(V10_PRED, persona_ids)

    mcnemar_summary = pd.DataFrame(
        [
            {
                "comparison": "claude_v1_vs_claude_v10",
                "n_paired": stats["n_paired"],
                "claude_v1_accuracy": stats["v1_accuracy"],
                "claude_v10_accuracy": stats["v10_accuracy"],
                "ate_pp_v10_minus_v1": stats["v10_accuracy"] - stats["v1_accuracy"],
                "v1_only_correct": stats["v1_only_correct"],
                "v10_only_correct": stats["v10_only_correct"],
                "both_correct": stats["both_correct"],
                "both_wrong": stats["both_wrong"],
                "mcnemar_p_value": stats["mcnemar_p_value"],
                "significant_at_0_05": stats["significant_at_0_05"],
            }
        ]
    )

    for out_dir in (OUT_DIR, LOCAL):
        if not out_dir:
            continue
        out_dir.mkdir(parents=True, exist_ok=True)
        xlsx_path = out_dir / XLSX_OUT
        with pd.ExcelWriter(xlsx_path, engine="openpyxl") as writer:
            v1_sheet.to_excel(writer, sheet_name="Claude_v1_120", index=False)
            v10_sheet.to_excel(writer, sheet_name="Claude_v10_120", index=False)
            mcnemar_summary.to_excel(writer, sheet_name="McNemar_Summary", index=False)
            pr_all.to_excel(writer, sheet_name="Confusion_Metrics", index=False)
            cm_v1.reset_index().to_excel(writer, sheet_name="Confusion_Claude_v1", index=False)
            cm_v10.reset_index().to_excel(writer, sheet_name="Confusion_Claude_v10", index=False)
        print(f"Wrote {xlsx_path}")

    build_docx(stats, cm_v1, cm_v10, pr_all)
    for out_dir in (OUT_DIR, LOCAL):
        print(f"Wrote {out_dir / DOCX_OUT}")


if __name__ == "__main__":
    main()
