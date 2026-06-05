"""McNemar comparison between two completed evaluation runs."""

from __future__ import annotations

import json
import os
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pandas as pd
from statsmodels.stats.contingency_tables import mcnemar

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def _load_predictions(csv_path: Path) -> pd.DataFrame:
    df = pd.read_csv(csv_path)
    if "persona_id" not in df.columns:
        raise ValueError(f"{csv_path.name} missing persona_id column")
    if "final_correct" in df.columns:
        ok_col = "final_correct"
    elif "label_match" in df.columns:
        ok_col = "label_match"
    else:
        raise ValueError(
            f"{csv_path.name} needs final_correct or label_match column"
        )
    out = df.drop_duplicates("persona_id", keep="last")[["persona_id", ok_col]].copy()
    out[ok_col] = out[ok_col].astype(bool)
    return out.rename(columns={ok_col: "correct"})


def _resolve_pred(path: Path | str, agent_root: Path | None = None) -> Path:
    p = Path(path)
    if p.is_file():
        return p
    root = agent_root or Path(__file__).resolve().parent.parent
    candidate = root / p
    if candidate.is_file():
        return candidate
    raise FileNotFoundError(f"predictions not found: {path}")


def compute_mcnemar(
    pred_a: Path | str,
    pred_b: Path | str,
    name_a: str,
    name_b: str,
    *,
    agent_root: Path | None = None,
) -> dict[str, Any]:
    pred_a = _resolve_pred(pred_a, agent_root)
    pred_b = _resolve_pred(pred_b, agent_root)
    a = _load_predictions(pred_a).rename(columns={"correct": "a_ok"})
    b = _load_predictions(pred_b).rename(columns={"correct": "b_ok"})
    m = a.merge(b, on="persona_id", how="inner")
    if m.empty:
        raise ValueError("no overlapping persona_id rows between the two runs")
    a_only = int((m["a_ok"] & ~m["b_ok"]).sum())
    b_only = int((~m["a_ok"] & m["b_ok"]).sum())
    both = int((m["a_ok"] & m["b_ok"]).sum())
    neither = int((~m["a_ok"] & ~m["b_ok"]).sum())
    table = [[both, b_only], [a_only, neither]]
    res = mcnemar(table, exact=True)
    acc_a = float(m["a_ok"].mean())
    acc_b = float(m["b_ok"].mean())
    return {
        "name_a": name_a,
        "name_b": name_b,
        "pred_a": str(pred_a),
        "pred_b": str(pred_b),
        "n_paired": len(m),
        "accuracy_a": acc_a,
        "accuracy_b": acc_b,
        "ate_b_minus_a_pp": (acc_b - acc_a) * 100.0,
        "a_only_correct": a_only,
        "b_only_correct": b_only,
        "both_correct": both,
        "both_wrong": neither,
        "discordant": a_only + b_only,
        "mcnemar_p_value": float(res.pvalue),
        "significant_at_0_05": bool(res.pvalue < 0.05),
        "pairs": m,
        "table": table,
        "created_utc": datetime.now(timezone.utc).isoformat(),
    }


def _pct(v: float) -> str:
    return f"{100.0 * v:.1f}%"


def write_comparison_outputs(stats: dict[str, Any], out_dir: Path) -> dict[str, Path]:
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    safe_a = "".join(c if c.isalnum() or c in "-_" else "_" for c in stats["name_a"])[:40]
    safe_b = "".join(c if c.isalnum() or c in "-_" else "_" for c in stats["name_b"])[:40]
    base = f"mcnemar_{safe_a}_vs_{safe_b}_{stamp}"

    pairs = stats.pop("pairs")
    json_path = out_dir / f"{base}.json"
    serial = {k: v for k, v in stats.items() if k != "table"}
    serial["mcnemar_table"] = {
        "rows": f"{stats['name_a']} correct",
        "cols": f"{stats['name_b']} correct",
        "cells": stats["table"],
    }
    json_path.write_text(json.dumps(serial, indent=2), encoding="utf-8")

    xlsx_path = out_dir / f"{base}.xlsx"
    summary = pd.DataFrame(
        [
            {"metric": "Group A", "value": stats["name_a"]},
            {"metric": "Group B", "value": stats["name_b"]},
            {"metric": "Paired personas", "value": stats["n_paired"]},
            {"metric": "Accuracy A", "value": _pct(stats["accuracy_a"])},
            {"metric": "Accuracy B", "value": _pct(stats["accuracy_b"])},
            {"metric": "ATE (B - A)", "value": f"{stats['ate_b_minus_a_pp']:+.1f} pp"},
            {"metric": "A only correct", "value": stats["a_only_correct"]},
            {"metric": "B only correct", "value": stats["b_only_correct"]},
            {"metric": "Both correct", "value": stats["both_correct"]},
            {"metric": "Both wrong", "value": stats["both_wrong"]},
            {"metric": "McNemar p-value", "value": f"{stats['mcnemar_p_value']:.4f}"},
            {
                "metric": "Significant at 0.05",
                "value": "Yes" if stats["significant_at_0_05"] else "No",
            },
        ]
    )
    with pd.ExcelWriter(xlsx_path, engine="openpyxl") as writer:
        summary.to_excel(writer, sheet_name="McNemar_Summary", index=False)
        pairs.to_excel(writer, sheet_name="Paired_Outcomes", index=False)

    md_path = out_dir / f"{base}.md"
    sig = "significant" if stats["significant_at_0_05"] else "not significant"
    md_path.write_text(
        "\n".join(
            [
                f"# McNemar: {stats['name_a']} vs {stats['name_b']}",
                "",
                f"- Paired personas: **{stats['n_paired']}**",
                f"- Accuracy A ({stats['name_a']}): **{_pct(stats['accuracy_a'])}**",
                f"- Accuracy B ({stats['name_b']}): **{_pct(stats['accuracy_b'])}**",
                f"- ATE (B − A): **{stats['ate_b_minus_a_pp']:+.1f} pp**",
                f"- Discordant: A-only={stats['a_only_correct']}, B-only={stats['b_only_correct']}",
                f"- McNemar exact p = **{stats['mcnemar_p_value']:.4f}** ({sig} at α=0.05)",
                "",
            ]
        ),
        encoding="utf-8",
    )

    docx_path = out_dir / f"{base}.docx"
    if HAS_DOCX:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        title = doc.add_heading(f"McNemar: {stats['name_a']} vs {stats['name_b']}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            f"Paired final-outcome accuracy on n={stats['n_paired']} personas. "
            f"{stats['name_a']}: {_pct(stats['accuracy_a'])}. "
            f"{stats['name_b']}: {_pct(stats['accuracy_b'])}. "
            f"ATE (B−A): {stats['ate_b_minus_a_pp']:+.1f} pp. "
            f"McNemar p={stats['mcnemar_p_value']:.4f} ({sig} at α=0.05)."
        )
        tbl = doc.add_table(rows=0, cols=2)
        for k, v in summary[["metric", "value"]].itertuples(index=False):
            row = tbl.add_row().cells
            row[0].text = str(k)
            row[1].text = str(v)
        _atomic_save_docx(doc, docx_path)

    stats["pairs"] = pairs
    return {"json": json_path, "xlsx": xlsx_path, "md": md_path, "docx": docx_path}


def _atomic_save_docx(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fd, tmp_name = tempfile.mkstemp(suffix=".docx", dir=str(path.parent))
    os.close(fd)
    tmp = Path(tmp_name)
    try:
        doc.save(str(tmp))
        if path.exists():
            path.unlink()
        tmp.replace(path)
    finally:
        if tmp.exists() and not path.exists():
            tmp.unlink(missing_ok=True)
