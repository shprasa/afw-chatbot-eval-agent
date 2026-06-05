"""Deliverables for OpenAI v1 vs Claude v1 comparison (120 personas).

Outputs:
  - powerbi_export/AFW_OpenAI_v1_vs_Claude_v1_Data.xlsx
    Sheets: OpenAI_v1_120, Claude_v1_120, McNemar_Summary
  - powerbi_export/AFW_OpenAI_v1_vs_Claude_v1_Report.docx
    McNemar, failure analysis + prompt recommendations (from rationale MD)
"""
from __future__ import annotations

import json
import os
import re
import tempfile
from math import comb
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from statsmodels.stats.contingency_tables import mcnemar

DESK = Path(__file__).resolve().parent
ARTIFACTS = DESK / "artifacts"
REPORTS = DESK / "reports"
SRC_120 = DESK / "AFW_120_User_Test_Cases_Original_Style_Short.xlsx"
OUT_DIR = DESK / "powerbi_export"
LOCAL = Path(os.environ.get("LOCALAPPDATA", "")) / "AFW_powerbi_export"

OPENAI_PRED = ARTIFACTS / "chatbot_live_predictions_original_style_short_openai.csv"
CLAUDE_PRED = ARTIFACTS / "chatbot_live_predictions_original_style_short_claude_v1.csv"
OPENAI_ACC = REPORTS / "chatbot_live_accuracy_original_style_short_openai.json"
CLAUDE_ACC = REPORTS / "chatbot_live_accuracy_original_style_short_claude_v1.json"

OPENAI_FAIL_MD = REPORTS / "chatbot_live_failure_analysis_original_style_short_openai.md"
OPENAI_PROMPT_MD = REPORTS / "chatbot_live_prompt_improvements_original_style_short_openai.md"
CLAUDE_FAIL_MD = REPORTS / "chatbot_live_failure_analysis_original_style_short_claude_v1.md"
CLAUDE_PROMPT_MD = REPORTS / "chatbot_live_prompt_improvements_original_style_short_claude_v1.md"

XLSX_OUT = "AFW_OpenAI_v1_vs_Claude_v1_Data.xlsx"
DOCX_OUT = "AFW_OpenAI_v1_vs_Claude_v1_Report.docx"

QN = [1, 2, 3, 4, 5, 6, 7, 8]
ROMAN = ["i", "ii", "iii", "iv", "v", "vi", "vii", "viii"]
CODE_TO_TEXT = {
    1: "eligible",
    0: "ineligible",
    2: "insufficient_information",
    3: "manual_review",
    "1": "eligible",
    "0": "ineligible",
    "2": "insufficient_information",
    "3": "manual_review",
}
CLASSES = ["eligible", "ineligible", "manual_review", "insufficient_information"]


def code_to_text(v) -> str:
    if pd.isna(v):
        return ""
    if isinstance(v, str) and not v.replace(".", "").isdigit():
        return v.strip()
    try:
        return CODE_TO_TEXT.get(int(float(v)), str(v))
    except (TypeError, ValueError):
        return str(v)


def build_120_sheet(pred_csv: Path) -> pd.DataFrame:
    src = pd.read_excel(SRC_120, sheet_name="120_test_cases", engine="openpyxl")
    src.columns = src.columns.str.lower()
    pred = pd.read_csv(pred_csv).drop_duplicates(subset="persona_id", keep="last")

    rows: list[dict] = []
    for _, s in src.iterrows():
        pid = s["persona_id"]
        p = pred[pred["persona_id"] == pid]
        p = p.iloc[0] if len(p) else None
        row: dict = {
            "persona_id": pid,
            "engineered_for (truth final)": s.get("engineered_for", ""),
        }
        for n, r in zip(QN, ROMAN):
            row[f"q{n} user input message"] = s.get(f"{r}. user input message", "")
            row[f"q{n} manual label"] = s.get(f"{r}. manual label", "")
            row[f"q{n} predicted"] = (
                code_to_text(p.get(f"pred_q{n}")) if p is not None else ""
            )
        truth_final = str(s.get("engineered_for", "")).strip()
        pred_final = code_to_text(p.get("predicted_label")) if p is not None else ""
        row["final manual label"] = truth_final
        row["final predicted"] = pred_final
        row["final match"] = (
            "" if not pred_final else ("TRUE" if truth_final == pred_final else "FALSE")
        )
        rows.append(row)
    return pd.DataFrame(rows)


def mcnemar_stats(openai_df: pd.DataFrame, claude_df: pd.DataFrame) -> dict:
    o = openai_df.drop_duplicates("persona_id")[["persona_id", "label_match"]].rename(
        columns={"label_match": "o_ok"}
    )
    c = claude_df.drop_duplicates("persona_id")[["persona_id", "label_match"]].rename(
        columns={"label_match": "c_ok"}
    )
    m = o.merge(c, on="persona_id")
    m["o_ok"] = m["o_ok"].astype(bool)
    m["c_ok"] = m["c_ok"].astype(bool)
    o_only = int((m["o_ok"] & ~m["c_ok"]).sum())
    c_only = int((~m["o_ok"] & m["c_ok"]).sum())
    both = int((m["o_ok"] & m["c_ok"]).sum())
    neither = int((~m["o_ok"] & ~m["c_ok"]).sum())
    table = [[both, c_only], [o_only, neither]]
    res = mcnemar(table, exact=True)
    n_disc = o_only + c_only
    return {
        "n_paired": len(m),
        "openai_accuracy": float(m["o_ok"].mean()),
        "claude_accuracy": float(m["c_ok"].mean()),
        "openai_only_correct": o_only,
        "claude_only_correct": c_only,
        "both_correct": both,
        "both_wrong": neither,
        "discordant": n_disc,
        "mcnemar_p_value": float(res.pvalue),
        "significant_at_0_05": bool(res.pvalue < 0.05),
        "pairs": m,
    }


def pct(v: float) -> str:
    return f"{100.0 * v:.1f}%"


def md_to_docx_sections(doc: Document, md_text: str, max_heading_level: int = 3) -> None:
    """Convert markdown headings and bullets into docx (subset)."""
    for line in md_text.splitlines():
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("---"):
            continue
        else:
            doc.add_paragraph(line.strip())


def atomic_save_docx(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fd, tmp_name = tempfile.mkstemp(suffix=".docx", dir=str(path.parent))
    os.close(fd)
    tmp = Path(tmp_name)
    try:
        doc.save(str(tmp))
        if path.exists():
            path.unlink()
        tmp.replace(path)
    finally:
        if tmp.exists() and not path.exists():
            tmp.unlink(missing_ok=True)


def build_docx(stats: dict) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading("AFW Chatbot Evaluation: OpenAI v1 vs Claude v1", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Dataset: AFW_120_User_Test_Cases_Original_Style_Short (F001-F120) | "
        "Prompt: System Prompt v1 | Cohort: 30 personas per final-outcome class"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Executive summary", level=1)
    sig = "statistically significant" if stats["significant_at_0_05"] else "not statistically significant"
    doc.add_paragraph(
        f"On paired final-outcome labels (n={stats['n_paired']}), OpenAI v1 achieved "
        f"{pct(stats['openai_accuracy'])} accuracy ({int(stats['openai_accuracy'] * stats['n_paired'])}/120) "
        f"and Claude v1 achieved {pct(stats['claude_accuracy'])} "
        f"({int(stats['claude_accuracy'] * stats['n_paired'])}/120). "
        f"The difference favors Claude by "
        f"{(stats['claude_accuracy'] - stats['openai_accuracy']) * 100:+.1f} percentage points. "
        f"McNemar's exact test on discordant pairs (OpenAI-only correct={stats['openai_only_correct']}, "
        f"Claude-only correct={stats['claude_only_correct']}) yields p={stats['mcnemar_p_value']:.4f}, "
        f"which is {sig} at alpha=0.05."
    )
    doc.add_paragraph(
        "Primary driver: OpenAI over-escalates insufficient_information personas to manual_review "
        "(15 cases vs Claude's 5), while Claude better holds incomplete screenings as insufficient_information."
    )

    doc.add_heading("2. McNemar test (OpenAI v1 vs Claude v1)", level=1)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"
    rows_kv = [
        ("Paired personas", str(stats["n_paired"])),
        ("OpenAI v1 accuracy", pct(stats["openai_accuracy"])),
        ("Claude v1 accuracy", pct(stats["claude_accuracy"])),
        ("Difference (Claude - OpenAI)", f"{(stats['claude_accuracy'] - stats['openai_accuracy']) * 100:+.1f} pp"),
        ("Both correct", str(stats["both_correct"])),
        ("OpenAI only correct", str(stats["openai_only_correct"])),
        ("Claude only correct", str(stats["claude_only_correct"])),
        ("Both wrong", str(stats["both_wrong"])),
        ("Discordant pairs", str(stats["discordant"])),
        ("McNemar exact two-sided p-value", f"{stats['mcnemar_p_value']:.4f}"),
        ("Significant at alpha=0.05", "Yes" if stats["significant_at_0_05"] else "No"),
    ]
    for k, v in rows_kv:
        cells = tbl.add_row().cells
        cells[0].text = k
        cells[1].text = v

    doc.add_paragraph(
        "Contingency (rows=OpenAI correct, cols=Claude correct): "
        f"[[{stats['both_correct']}, {stats['claude_only_correct']}], "
        f"[{stats['openai_only_correct']}, {stats['both_wrong']}]]"
    )

    doc.add_heading("3. Headline accuracy by model", level=1)
    for label, acc_path in [("OpenAI v1", OPENAI_ACC), ("Claude v1", CLAUDE_ACC)]:
        acc = json.loads(acc_path.read_text(encoding="utf-8"))
        doc.add_heading(label, level=2)
        doc.add_paragraph(
            f"Final outcome accuracy: {pct(acc['final_outcome_accuracy'])} "
            f"({acc['correct']}/{acc['n_scored']}). "
            f"Checkpoint-level accuracy: {pct(acc.get('question_level_accuracy', 0))}."
        )
        by_class = acc.get("by_truth_class", {})
        for cls in CLASSES:
            key = f"is_{cls}"
            info = by_class.get(key, {})
            doc.add_paragraph(
                f"{cls} recall: {pct(info.get('recall', 0))} ({info.get('correct', 0)}/{info.get('n_in_class', 0)})",
                style="List Bullet",
            )

    doc.add_heading("4. Failure analysis — OpenAI v1 (rationale-driven)", level=1)
    doc.add_paragraph(
        "Source: live eval transcripts; API sessionData fields (outcome_notes, *_rationale, "
        "outcome_reason_codes). Full detail in appendix markdown."
    )
    if OPENAI_FAIL_MD.is_file():
        md_to_docx_sections(doc, OPENAI_FAIL_MD.read_text(encoding="utf-8"))

    doc.add_heading("5. Failure analysis — Claude v1 (rationale-driven)", level=1)
    if CLAUDE_FAIL_MD.is_file():
        md_to_docx_sections(doc, CLAUDE_FAIL_MD.read_text(encoding="utf-8"))

    doc.add_heading("6. System prompt recommendations — OpenAI v1", level=1)
    doc.add_paragraph(
        "Edits below are derived from OpenAI v1 error patterns and API rationale on failures. "
        "Apply to artifacts/prompt_extract_v1.txt (System Prompt v1)."
    )
    if OPENAI_PROMPT_MD.is_file():
        md_to_docx_sections(doc, OPENAI_PROMPT_MD.read_text(encoding="utf-8"))

    doc.add_heading("7. System prompt recommendations — Claude v1", level=1)
    if CLAUDE_PROMPT_MD.is_file():
        md_to_docx_sections(doc, CLAUDE_PROMPT_MD.read_text(encoding="utf-8"))

    doc.add_heading("8. Shared prompt v1 recommendations (both backends)", level=1)
    doc.add_paragraph(
        "Both models share the same System Prompt v1. Prioritize these cross-cutting edits:"
    )
    for line in [
        "Stop escalating missing required fields to manual_review; use insufficient_information + DATA_INCOMPLETE.",
        "When all required fields are present but judgment is needed, use manual_review with criterion-specific rationales.",
        "Commit to ineligible when a non-negotiable disqualifier is confirmed.",
        "Gate eligible on all checkpoints satisfied and no active manual-review trigger.",
        "Require per-turn *_rationale refresh before each eligibility_outcome update.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("9. Data companion", level=1)
    doc.add_paragraph(
        f"Wide-format predictions workbook: {XLSX_OUT} "
        f"(sheets OpenAI_v1_120 and Claude_v1_120)."
    )

    for out_dir in (OUT_DIR, LOCAL):
        if out_dir:
            atomic_save_docx(doc, out_dir / DOCX_OUT)


def main() -> None:
    openai_sheet = build_120_sheet(OPENAI_PRED)
    claude_sheet = build_120_sheet(CLAUDE_PRED)
    o_df = pd.read_csv(OPENAI_PRED)
    c_df = pd.read_csv(CLAUDE_PRED)
    stats = mcnemar_stats(o_df, c_df)

    mcnemar_summary = pd.DataFrame(
        [
            {
                "comparison": "openai_v1_vs_claude_v1",
                "n_paired": stats["n_paired"],
                "openai_v1_accuracy": stats["openai_accuracy"],
                "claude_v1_accuracy": stats["claude_accuracy"],
                "ate_pp_claude_minus_openai": stats["claude_accuracy"] - stats["openai_accuracy"],
                "openai_only_correct": stats["openai_only_correct"],
                "claude_only_correct": stats["claude_only_correct"],
                "both_correct": stats["both_correct"],
                "both_wrong": stats["both_wrong"],
                "mcnemar_p_value": stats["mcnemar_p_value"],
                "significant_at_0_05": stats["significant_at_0_05"],
            }
        ]
    )

    for out_dir in (OUT_DIR, LOCAL):
        if not out_dir:
            continue
        out_dir.mkdir(parents=True, exist_ok=True)
        xlsx_path = out_dir / XLSX_OUT
        with pd.ExcelWriter(xlsx_path, engine="openpyxl") as writer:
            openai_sheet.to_excel(writer, sheet_name="OpenAI_v1_120", index=False)
            claude_sheet.to_excel(writer, sheet_name="Claude_v1_120", index=False)
            mcnemar_summary.to_excel(writer, sheet_name="McNemar_Summary", index=False)
        print(f"Wrote {xlsx_path}")

    build_docx(stats)
    for out_dir in (OUT_DIR, LOCAL):
        print(f"Wrote {out_dir / DOCX_OUT}")


if __name__ == "__main__":
    main()
