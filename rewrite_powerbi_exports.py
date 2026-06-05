"""
Full Power BI export: personas, turns (user messages), checkpoints, API rationales, McNemar.
Run: python rewrite_powerbi_exports.py
"""
from __future__ import annotations

import json
import os
import tempfile
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from statsmodels.stats.contingency_tables import mcnemar

DESK = Path(__file__).resolve().parent
ARTIFACTS = DESK / "artifacts"
TRUTH_XLSX = DESK / "AFW_120_User_Test_Cases_Original_Style_Short.xlsx"
COHORT = "120_OriginalStyleShort"

OUT_DIRS = [
    DESK / "powerbi_export",
    Path(os.environ.get("LOCALAPPDATA", "")) / "AFW_powerbi_export",
]
XLSX_NAME = "AFW_Eval_PowerBI_Data.xlsx"
DOCX_NAME = "AFW_PowerBI_Guide.docx"

CODE_TO_LABEL = {
    0: "ineligible",
    1: "eligible",
    2: "insufficient_information",
    3: "manual_review",
}
LABEL_TO_CODE = {v: k for k, v in CODE_TO_LABEL.items()}
QUESTIONS = [f"q{i}" for i in range(1, 9)]
ROMAN = ["i", "ii", "iii", "iv", "v", "vi", "vii", "viii"]
TURN_ORDER = ["opening"] + QUESTIONS

NOTE_FIELDS = (
    "outcome_notes",
    "summary_notes",
    "summary_compelling_or_financial_need_notes",
    "summary_sentiment_analysis",
    "manual_review_rationale",
)

ARMS: list[tuple[str, str, str]] = [
    (
        "openai_v1",
        "chatbot_live_predictions_original_style_short_openai.csv",
        "chatbot_live_transcripts_original_style_short_openai.jsonl",
    ),
    (
        "openai_v10",
        "chatbot_live_predictions_original_style_short_openai_v10.csv",
        "chatbot_live_transcripts_original_style_short_openai_v10.jsonl",
    ),
    (
        "claude_v1",
        "chatbot_live_predictions_original_style_short_claude_v1.csv",
        "chatbot_live_transcripts_original_style_short_claude_v1.jsonl",
    ),
    (
        "claude_v10",
        "chatbot_live_predictions_original_style_short_claude.csv",
        "chatbot_live_transcripts_original_style_short_claude.jsonl",
    ),
]


def code_to_label(code: Any) -> str:
    if code is None or (isinstance(code, float) and np.isnan(code)):
        return ""
    try:
        c = int(code)
        return CODE_TO_LABEL.get(c, str(code))
    except (TypeError, ValueError):
        return str(code).strip()


def extract_rationales_from_session(sess: dict[str, Any]) -> list[dict[str, str]]:
    """Long-format rationale rows from sessionData."""
    if not isinstance(sess, dict):
        return []
    rows: list[dict[str, str]] = []
    for key, val in sess.items():
        if not isinstance(key, str) or val is None:
            continue
        text = str(val).strip()
        if not text:
            continue
        if key.endswith("_rationale") or key in NOTE_FIELDS or key == "outcome_reason_codes":
            if key == "outcome_reason_codes":
                if isinstance(val, list):
                    text = ", ".join(str(x) for x in val)
                else:
                    text = str(val)
            rows.append({"field_name": key, "field_value": text})
    elig = sess.get("eligibility_outcome")
    if elig is not None and str(elig).strip():
        rows.append({"field_name": "eligibility_outcome", "field_value": str(elig).strip()})
    return rows


def load_truth_workbook() -> pd.DataFrame:
    df = pd.read_excel(TRUTH_XLSX, sheet_name="120_test_cases", engine="openpyxl")
    df.columns = df.columns.str.lower()
    return df.set_index("persona_id")


def load_transcripts(path: Path) -> dict[str, dict]:
    by_pid: dict[str, dict] = {}
    if not path.is_file():
        return by_pid
    for line in path.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        try:
            rec = json.loads(line)
        except json.JSONDecodeError:
            continue
        pid = str(rec.get("persona_id", "")).strip()
        if pid:
            by_pid[pid] = rec
    return by_pid


def _cell_str(v: Any, max_len: int = 8000) -> str:
    if v is None:
        return ""
    s = str(v).strip()
    if len(s) > max_len:
        return s[: max_len - 3] + "..."
    return s


def build_arm_tables(
    arm_id: str, pred_csv: Path, transcript_jsonl: Path, truth_df: pd.DataFrame
) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    transcripts = load_transcripts(transcript_jsonl)
    pred_df = pd.read_csv(pred_csv).drop_duplicates("persona_id", keep="last")

    persona_rows: list[dict] = []
    turn_rows: list[dict] = []
    rationale_rows: list[dict] = []

    for _, prow in pred_df.iterrows():
        pid = str(prow["persona_id"])
        rec = transcripts.get(pid, {})
        score = rec.get("score") or {}
        session_id = str(rec.get("session_id") or prow.get("session_id") or "")
        per_q = score.get("per_question") or {}

        truth_row = truth_df.loc[pid] if pid in truth_df.index else None
        engineered_for = ""
        if truth_row is not None:
            engineered_for = str(truth_row.get("engineered_for", "") or "").strip()

        truth_final_code = _to_int(prow.get("truth_code"))
        pred_final_code = _to_int(prow.get("predicted_code"))
        final_sess = rec.get("final_session_data") or {}
        turn_list = rec.get("turns") or []

        was_completed = final_sess.get("was_session_completed")
        if was_completed is None:
            screening_completed = int(
                pred_final_code is not None
                and not rec.get("error")
                and len(turn_list) >= 8
            )
        else:
            screening_completed = int(bool(was_completed))

        persona_rows.append(
            {
                "persona_id": pid,
                "arm_id": arm_id,
                "cohort_id": COHORT,
                "session_id": session_id,
                "engineered_for": engineered_for,
                "truth_final_code": truth_final_code,
                "truth_final_label": code_to_label(truth_final_code),
                "pred_final_code": pred_final_code,
                "pred_final_label": code_to_label(pred_final_code),
                "final_correct": _to_int(prow.get("label_match"), bool_ok=True),
                "screening_completed": screening_completed,
                "final_eligibility_outcome": _cell_str(
                    final_sess.get("eligibility_outcome"), 200
                ),
                "final_reason_codes": _cell_str(
                    final_sess.get("outcome_reason_codes"), 500
                ),
                "manual_review_flag": int(bool(final_sess.get("manual_review_flag"))),
                "fn_missed_eligible": int(
                    truth_final_code == 1 and pred_final_code != 1
                ),
                "fn_missed_manual_review": int(
                    truth_final_code == 3 and pred_final_code != 3
                ),
                "fn_missed_ineligible": int(
                    truth_final_code == 0 and pred_final_code != 0
                ),
                "unnecessary_manual_review": int(
                    truth_final_code == 2 and pred_final_code == 3
                ),
                "unsafe_approve": int(
                    truth_final_code in (0, 3) and pred_final_code == 1
                ),
                "error": _cell_str(rec.get("error"), 500),
            }
        )

        for scope, sess in [("final", final_sess)]:
            for r in extract_rationales_from_session(sess):
                rationale_rows.append(
                    {
                        "persona_id": pid,
                        "arm_id": arm_id,
                        "cohort_id": COHORT,
                        "scope": scope,
                        "question": "",
                        "field_name": r["field_name"],
                        "field_value": r["field_value"],
                    }
                )

        for key, val in final_sess.items():
            if isinstance(key, str) and key.endswith("_flag"):
                rationale_rows.append(
                    {
                        "persona_id": pid,
                        "arm_id": arm_id,
                        "cohort_id": COHORT,
                        "scope": "criterion_flag",
                        "question": "",
                        "field_name": key,
                        "field_value": _cell_str(val, 100),
                    }
                )

        for t_idx, turn in enumerate(turn_list):
            q = str(turn.get("question", "") or "")
            pq = per_q.get(q, {}) if q in per_q else {}
            truth_code = _to_int(pq.get("truth_code"))
            pred_code = _to_int(pq.get("pred_code"))
            if pred_code is None:
                pred_code = _label_to_code(turn.get("predicted_outcome"))

            truth_manual = ""
            if truth_row is not None and q in QUESTIONS:
                roman = ROMAN[QUESTIONS.index(q)]
                col = f"{roman}. manual label"
                if col in truth_row.index:
                    truth_manual = _cell_str(truth_row.get(col), 200)
            elif truth_row is not None and q == "opening":
                truth_manual = _cell_str(truth_row.get("simulated_user_message"), 200)

            turn_rows.append(
                {
                    "persona_id": pid,
                    "arm_id": arm_id,
                    "cohort_id": COHORT,
                    "session_id": session_id,
                    "turn_index": t_idx,
                    "question": q,
                    "input_col": _cell_str(turn.get("input_col"), 120),
                    "user_message": _cell_str(turn.get("user_message"), 8000),
                    "assistant_response": _cell_str(turn.get("assistant_response"), 8000),
                    "truth_checkpoint_code": truth_code,
                    "truth_checkpoint_label": code_to_label(truth_code),
                    "truth_manual_label": truth_manual,
                    "pred_checkpoint_code": pred_code,
                    "pred_checkpoint_label": code_to_label(pred_code),
                    "checkpoint_correct": _to_int(pq.get("match"), bool_ok=True)
                    if pq
                    else None,
                    "session_eligibility_outcome": _cell_str(
                        turn.get("session_eligibility_outcome"), 200
                    ),
                    "outcome_reason_codes": _cell_str(
                        turn.get("outcome_reason_codes"), 500
                    ),
                    "is_manual_review_flag": int(bool(turn.get("isManualReview"))),
                }
            )

            sess_turn = turn.get("session_data") or turn.get("session_rationales")
            if isinstance(sess_turn, dict) and sess_turn:
                for r in extract_rationales_from_session(sess_turn):
                    rationale_rows.append(
                        {
                            "persona_id": pid,
                            "arm_id": arm_id,
                            "cohort_id": COHORT,
                            "scope": "turn",
                            "question": q,
                            "field_name": r["field_name"],
                            "field_value": r["field_value"],
                        }
                    )

    fact_persona = pd.DataFrame(persona_rows)
    fact_turn = pd.DataFrame(turn_rows)
    fact_rationale = pd.DataFrame(rationale_rows)
    fact_checkpoint = fact_turn[
        [
            "persona_id",
            "arm_id",
            "cohort_id",
            "question",
            "turn_index",
            "truth_checkpoint_code",
            "truth_checkpoint_label",
            "truth_manual_label",
            "pred_checkpoint_code",
            "pred_checkpoint_label",
            "checkpoint_correct",
        ]
    ].copy()
    return fact_persona, fact_turn, fact_checkpoint, fact_rationale


def _label_to_code(label: Any) -> int | None:
    if label is None:
        return None
    s = str(label).strip().lower()
    return LABEL_TO_CODE.get(s)


def _to_int(v: Any, bool_ok: bool = False) -> int | None:
    if v is None or (isinstance(v, float) and np.isnan(v)):
        return None
    if bool_ok:
        return int(bool(v))
    try:
        return int(v)
    except (TypeError, ValueError):
        return None


def mcnemar_stats(df_a: pd.DataFrame, df_b: pd.DataFrame, comparison_id: str) -> tuple:
    a = df_a.set_index("persona_id")
    b = df_b.set_index("persona_id")
    common = sorted(set(a.index) & set(b.index))
    ya = np.array([int(a.loc[p, "final_correct"]) for p in common])
    yb = np.array([int(b.loc[p, "final_correct"]) for p in common])
    both_ok = int(np.sum(ya & yb))
    both_bad = int(np.sum((1 - ya) & (1 - yb)))
    bw = int(np.sum((1 - ya) & yb))
    aw = int(np.sum(ya & (1 - yb)))
    res = mcnemar([[both_ok, aw], [bw, both_bad]], exact=True, correction=True)
    stats = {
        "comparison_id": comparison_id,
        "n_paired": len(common),
        "arm_a_correct": int(ya.sum()),
        "arm_b_correct": int(yb.sum()),
        "arm_a_accuracy": round(float(ya.mean()), 6),
        "arm_b_accuracy": round(float(yb.mean()), 6),
        "ate_pp_arm_b_minus_a": round((float(yb.mean()) - float(ya.mean())) * 100, 4),
        "both_correct": both_ok,
        "both_wrong": both_bad,
        "arm_a_wrong_arm_b_right": bw,
        "arm_a_right_arm_b_wrong": aw,
        "mcnemar_statistic": float(res.statistic) if res.statistic is not None else None,
        "mcnemar_p_value": float(res.pvalue),
        "significant_alpha_0_05": int(res.pvalue < 0.05),
        "cohort_id": COHORT,
    }
    pairs = []
    for p in common:
        ca, cb = int(a.loc[p, "final_correct"]), int(b.loc[p, "final_correct"])
        pairs.append(
            {
                "comparison_id": comparison_id,
                "persona_id": p,
                "cohort_id": COHORT,
                "arm_a_correct": ca,
                "arm_b_correct": cb,
                "discordant": int(ca != cb),
                "arm_a_pred_code": _to_int(a.loc[p, "pred_final_code"]),
                "arm_b_pred_code": _to_int(b.loc[p, "pred_final_code"]),
                "truth_final_code": _to_int(a.loc[p, "truth_final_code"]),
                "arm_a_pred_label": a.loc[p, "pred_final_label"],
                "arm_b_pred_label": b.loc[p, "pred_final_label"],
                "truth_final_label": a.loc[p, "truth_final_label"],
            }
        )
    return pd.DataFrame(pairs), stats


def summary_arm(df: pd.DataFrame) -> dict:
    scored = df.dropna(subset=["final_correct"])
    n = len(scored)
    c = int(scored["final_correct"].sum()) if n else 0
    return {
        "arm_id": df["arm_id"].iloc[0],
        "cohort_id": COHORT,
        "n_personas": len(df),
        "n_scored": n,
        "n_correct": c,
        "accuracy": round(c / n, 6) if n else None,
    }


def recall_by_class(df: pd.DataFrame) -> pd.DataFrame:
    out = []
    for code in CODE_TO_LABEL:
        sub = df[df["truth_final_code"] == code].dropna(subset=["final_correct"])
        if len(sub) == 0:
            continue
        out.append(
            {
                "arm_id": df["arm_id"].iloc[0],
                "cohort_id": COHORT,
                "truth_class_code": code,
                "n_in_class": len(sub),
                "n_correct": int(sub["final_correct"].sum()),
                "recall": round(float(sub["final_correct"].mean()), 6),
            }
        )
    return pd.DataFrame(out)


def write_xlsx(path: Path, sheets: dict[str, pd.DataFrame]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fd, tmp_name = tempfile.mkstemp(suffix=".xlsx", dir=str(path.parent))
    os.close(fd)
    tmp = Path(tmp_name)
    try:
        wb = Workbook()
        ws0 = wb.active
        ws0.title = "README"
        first = True
        for title, df in sheets.items():
            if first:
                for row in dataframe_to_rows(df, index=False, header=True):
                    ws0.append(row)
                first = False
            else:
                name = title[:31]
                ws = wb.create_sheet(title=name)
                for row in dataframe_to_rows(df, index=False, header=True):
                    ws.append(row)
        wb.save(tmp)
        wb.close()
        if path.exists():
            path.unlink()
        tmp.replace(path)
    finally:
        if tmp.exists() and not path.exists():
            tmp.unlink(missing_ok=True)


def template_eval_screenings() -> pd.DataFrame:
    """Empty template for gold evaluation screenings (fill when live data arrives)."""
    return pd.DataFrame(
        columns=[
            "screening_id",
            "persona_or_session_key",
            "source_type",
            "eval_run_date",
            "arm_id",
            "model_host",
            "prompt_version",
            "truth_final_code",
            "truth_final_label",
            "pred_final_code",
            "pred_final_label",
            "screening_completed",
            "turn_count",
            "notes",
        ]
    )


def template_actual_chatbot_screenings() -> pd.DataFrame:
    """Empty template for production chatbot screenings (fill from ops export)."""
    return pd.DataFrame(
        columns=[
            "screening_id",
            "live_session_id",
            "session_start_utc",
            "production_arm_label",
            "pred_final_code",
            "pred_final_label",
            "screening_completed",
            "referred_to_coordinator",
            "coordinator_touch_required",
            "outcome_reason_codes",
            "notes",
        ]
    )


def _add_bullets(doc: Any, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def write_docx(path: Path, mcnemar_df: pd.DataFrame) -> None:
    from docx import Document
    from docx.shared import Pt

    path.parent.mkdir(parents=True, exist_ok=True)
    fd, tmp_name = tempfile.mkstemp(suffix=".docx", dir=str(path.parent))
    os.close(fd)
    tmp = Path(tmp_name)
    try:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        doc.add_heading(
            "AFW New Horizons Initiative - Power BI Dashboard Guide", 0
        )
        doc.add_paragraph(
            "This guide maps AFW_Eval_PowerBI_Data.xlsx to New Horizons reporting: "
            "reduce false negatives, reduce coordinator manual interventions, "
            "track completed screenings by eligibility class, and compare eval vs production."
        )

        doc.add_heading("1. Import and data model", level=1)
        _add_bullets(
            doc,
            [
                "Get Data > Excel > AFW_Eval_PowerBI_Data.xlsx > load all sheets.",
                "Prefer the copy in %LOCALAPPDATA%\\AFW_powerbi_export if OneDrive shows a broken file icon.",
                "Set all *_code columns to Whole Number; screening_completed and error flags to Whole Number (0/1).",
                "Relationships: dim_arm[arm_id] -> fact_persona, fact_turn, fact_rationale (many-to-one).",
                "fact_persona[persona_id, arm_id] -> fact_turn and fact_rationale (one-to-many).",
                "dim_encoding[code] -> optional label tooltips for truth/pred codes.",
                "Do not relate tpl_eval_screenings or tpl_actual_chatbot until production rows are added.",
            ],
        )

        doc.add_heading("2. Outcome encoding (required)", level=1)
        doc.add_paragraph(
            "All dashboards use numeric codes: 0=ineligible, 1=eligible, "
            "2=insufficient_information, 3=manual_review. "
            "Labels are in dim_encoding and fact_persona.*_label columns."
        )

        doc.add_heading("3. New Horizons goal 1 - Reduce false negatives", level=1)
        doc.add_paragraph(
            "False negative means the chatbot failed to surface the correct screening outcome "
            "relative to gold truth (evaluation personas). Use fact_persona columns:"
        )
        _add_bullets(
            doc,
            [
                "fn_missed_eligible (1): truth=eligible but pred is not eligible - passenger may be wrongly denied.",
                "fn_missed_manual_review (1): truth=manual_review but pred is not MR - coordinator may not be engaged when needed.",
                "fn_missed_ineligible (1): truth=ineligible but pred is not ineligible - disqualifier not committed.",
                "unsafe_approve (1): truth is ineligible or manual_review but pred=eligible - highest safety risk.",
            ],
        )
        doc.add_paragraph("Recommended visuals:")
        _add_bullets(
            doc,
            [
                "KPI cards: COUNTROWS where fn_missed_* = 1, sliced by arm_id; show rate = FN / COUNT(fact_persona).",
                "Matrix: rows=arm_id, columns=truth_final_label, values=sum of fn_missed_manual_review.",
                "Drill-through: fact_persona -> fact_turn (user_message) -> fact_rationale (scope=final, outcome_notes).",
                "Filter fact_rationale field_name contains 'rationale' to read API reasoning for each FN case.",
            ],
        )
        doc.add_paragraph("DAX examples:")
        for dax in [
            "FN Rate Missed MR = DIVIDE(CALCULATE(COUNTROWS(fact_persona), fact_persona[fn_missed_manual_review]=1), COUNTROWS(fact_persona))",
            "FN Count Any = CALCULATE(COUNTROWS(fact_persona), fact_persona[fn_missed_eligible]=1 || fact_persona[fn_missed_manual_review]=1 || fact_persona[fn_missed_ineligible]=1)",
            "Unsafe Approve Count = CALCULATE(COUNTROWS(fact_persona), fact_persona[unsafe_approve]=1)",
        ]:
            doc.add_paragraph(dax, style="List Bullet")

        doc.add_heading(
            "4. New Horizons goal 2 - Reduce manual coordinator interventions", level=1
        )
        doc.add_paragraph(
            "Proxy for unnecessary coordinator load: predicted manual_review when gold truth "
            "is insufficient_information (user still owes facts, not ready for coordinator)."
        )
        _add_bullets(
            doc,
            [
                "unnecessary_manual_review (1): truth=insufficient, pred=manual_review - over-escalation pattern.",
                "fn_missed_manual_review (1): under-escalation - opposite problem.",
                "Compare unnecessary_manual_review rate across arm_id to pick prompt/model with lowest over-escalation.",
            ],
        )
        doc.add_paragraph("DAX examples:")
        for dax in [
            "Unnecessary MR Rate = DIVIDE(CALCULATE(COUNTROWS(fact_persona), fact_persona[unnecessary_manual_review]=1), COUNTROWS(fact_persona))",
            "MR Pred Rate = DIVIDE(CALCULATE(COUNTROWS(fact_persona), fact_persona[pred_final_code]=3), COUNTROWS(fact_persona))",
        ]:
            doc.add_paragraph(dax, style="List Bullet")

        doc.add_heading(
            "5. Completed screenings by eligibility class", level=1
        )
        doc.add_paragraph(
            "A screening is completed when screening_completed=1 on fact_persona "
            "(from API was_session_completed, or proxy: 8+ turns and no error)."
        )
        doc.add_paragraph(
            "Segment completed screenings by truth eligibility class (gold) and by predicted class:"
        )
        _add_bullets(
            doc,
            [
                "Stacked bar 100%: Axis=arm_id, Legend=truth_final_label, Filter screening_completed=1.",
                "Second chart: Legend=pred_final_label for model outcome mix on completed only.",
                "Table: arm_id x truth_final_label with count and % of completed screenings.",
            ],
        )
        doc.add_paragraph("DAX examples:")
        for dax in [
            "Completed Screenings = CALCULATE(COUNTROWS(fact_persona), fact_persona[screening_completed]=1)",
            "Pct Completed (by arm) = DIVIDE([Completed Screenings], CALCULATE(COUNTROWS(fact_persona), ALLEXCEPT(fact_persona, fact_persona[arm_id])))",
            "Completed by Truth Class = CALCULATE(COUNTROWS(fact_persona), fact_persona[screening_completed]=1, VALUES(fact_persona[truth_final_label]))",
        ]:
            doc.add_paragraph(dax, style="List Bullet")

        doc.add_heading(
            "6. Screening outcome mix (what % ended as 0/1/2/3)", level=1
        )
        doc.add_paragraph(
            "Of all screenings (or completed-only subset), show distribution of pred_final_code "
            "by arm_id (openai_v1, openai_v10, claude_v1, claude_v10):"
        )
        _add_bullets(
            doc,
            [
                "100% stacked column: arm_id on axis, pred_final_code or pred_final_label as legend.",
                "Donut per arm: share of outcomes 0/1/2/3.",
                "Add slicer: screening_completed (All / Completed only / Not completed).",
            ],
        )
        doc.add_paragraph("DAX:")
        doc.add_paragraph(
            "Outcome Share = DIVIDE(COUNTROWS(fact_persona), CALCULATE(COUNTROWS(fact_persona), ALLEXCEPT(fact_persona, fact_persona[arm_id], fact_persona[pred_final_code])))",
            style="List Bullet",
        )

        doc.add_heading(
            "7. Per eligibility criteria outcome (completed screenings)", level=1
        )
        doc.add_paragraph(
            "Criterion-level flags from API sessionData are in fact_rationale where "
            "scope=criterion_flag (e.g. service_region_flag, meets_minimum_notice_flag). "
            "Values are typically true/false/blank from the API."
        )
        _add_bullets(
            doc,
            [
                "Filter fact_persona screening_completed=1, relate to fact_rationale on persona_id+arm_id.",
                "Matrix: field_name (criterion) x field_value, count of personas.",
                "For pass/fail style: treat 'true'/'True' as pass, 'false' as fail, blank as not assessed.",
                "Pair with *_rationale text in scope=final for explanation.",
            ],
        )
        doc.add_paragraph(
            "Criteria fields to monitor (in fact_rationale, scope=criterion_flag): "
            "service_region_flag, within_distance_limit_for_small_plane_flag, "
            "meets_minimum_notice_flag, is_mobile, is_medically_stable, "
            "has_financial_or_compelling_need, dv_context_applicable."
        )

        doc.add_heading(
            "8. Evaluation vs actual chatbot screenings (templates)", level=1
        )
        doc.add_paragraph(
            "Two empty template sheets are included for future production comparison. "
            "No live production data is loaded yet."
        )
        doc.add_paragraph("Sheet tpl_eval_screenings - fill from evaluation runs:")
        _add_bullets(
            doc,
            [
                "screening_id, persona_or_session_key, source_type=evaluation, eval_run_date, arm_id, model_host, prompt_version.",
                "truth_final_code/label, pred_final_code/label, screening_completed, turn_count, notes.",
            ],
        )
        doc.add_paragraph("Sheet tpl_actual_chatbot - fill from production exports:")
        _add_bullets(
            doc,
            [
                "screening_id, live_session_id, session_start_utc, production_arm_label.",
                "pred_final_code/label, screening_completed, referred_to_coordinator, coordinator_touch_required, outcome_reason_codes.",
            ],
        )
        doc.add_paragraph(
            "When production data exists: append tpl_actual_chatbot to a new fact_actual_screenings table, "
            "union with fact_persona (rename columns), build a source_type column (evaluation vs production), "
            "and duplicate pages 3-7 with a source_type slicer."
        )

        doc.add_heading("9. Suggested dashboard pages (8 pages)", level=1)
        pages = [
            ("Executive", "KPIs: accuracy by arm, completed %, FN rates, unnecessary MR rate, McNemar p-values from mcnemar_results."),
            ("False negatives", "FN matrix, drill persona list, rationale snippets for top FN patterns."),
            ("Coordinator load", "unnecessary_manual_review vs fn_missed_manual_review by arm."),
            ("Completion", "screening_completed rate by arm and truth class."),
            ("Outcome mix", "100% stacked pred outcome 0/1/2/3 by arm."),
            ("Criteria", "criterion_flag matrix for completed screenings only."),
            ("Chat detail", "fact_turn table with user_message / assistant_response for selected persona."),
            ("Eval vs production", "placeholder page using tpl_* sheets until live data wired."),
        ]
        for title, body in pages:
            doc.add_heading(title, level=2)
            doc.add_paragraph(body)

        doc.add_heading("10. Key tables reference", level=1)
        _add_bullets(
            doc,
            [
                "fact_persona: one row per persona per arm; primary grain for outcomes and FN flags.",
                "fact_turn: chat transcript turns with user_message and assistant_response.",
                "fact_rationale: long-format API rationales (scope=final, turn, criterion_flag).",
                "fact_checkpoint: checkpoint truth vs pred per question.",
                "fact_summary_arm / fact_recall_by_class: pre-aggregated accuracy.",
                "mcnemar_results + mcnemar_* pair sheets: statistical comparisons.",
            ],
        )

        doc.add_heading("11. McNemar results (reference)", level=1)
        tbl = doc.add_table(rows=1, cols=6)
        for i, h in enumerate(
            ["comparison", "n", "acc_a", "acc_b", "ate_pp", "p_value"]
        ):
            tbl.rows[0].cells[i].text = h
        for _, r in mcnemar_df.iterrows():
            cells = tbl.add_row().cells
            cells[0].text = str(r["comparison_id"])
            cells[1].text = str(int(r["n_paired"]))
            cells[2].text = f"{float(r['arm_a_accuracy']):.4f}"
            cells[3].text = f"{float(r['arm_b_accuracy']):.4f}"
            cells[4].text = f"{float(r['ate_pp_arm_b_minus_a']):.2f}"
            cells[5].text = f"{float(r['mcnemar_p_value']):.4f}"

        doc.save(str(tmp))
        if path.exists():
            path.unlink()
        tmp.replace(path)
    finally:
        if tmp.exists() and not path.exists():
            tmp.unlink(missing_ok=True)


def main() -> None:
    for d in OUT_DIRS:
        if d.is_dir():
            for f in d.glob("*.xlsx"):
                f.unlink()
            for f in d.glob("*.docx"):
                f.unlink()

    truth_df = load_truth_workbook()

    all_persona: list[pd.DataFrame] = []
    all_turn: list[pd.DataFrame] = []
    all_checkpoint: list[pd.DataFrame] = []
    all_rationale: list[pd.DataFrame] = []
    personas_by_arm: dict[str, pd.DataFrame] = {}

    for arm_id, pred_name, jsonl_name in ARMS:
        pred_path = ARTIFACTS / pred_name
        jsonl_path = ARTIFACTS / jsonl_name
        if not pred_path.is_file():
            print("SKIP missing", arm_id, pred_path)
            continue
        fp, ft, fc, fr = build_arm_tables(arm_id, pred_path, jsonl_path, truth_df)
        all_persona.append(fp)
        all_turn.append(ft)
        all_checkpoint.append(fc)
        all_rationale.append(fr)
        personas_by_arm[arm_id] = fp
        print(f"{arm_id}: personas={len(fp)} turns={len(ft)} rationales={len(fr)}")

    fact_persona = pd.concat(all_persona, ignore_index=True)
    fact_turn = pd.concat(all_turn, ignore_index=True)
    fact_checkpoint = pd.concat(all_checkpoint, ignore_index=True)
    fact_rationale = pd.concat(all_rationale, ignore_index=True)

    dim_encoding = pd.DataFrame([{"code": k, "label": v} for k, v in sorted(CODE_TO_LABEL.items())])
    dim_arm = pd.DataFrame([{"arm_id": a[0], "pred_file": a[1], "transcript_file": a[2]} for a in ARMS])
    dim_cohort = pd.DataFrame(
        [{"cohort_id": COHORT, "n_personas": 120, "persona_range": "F001-F120"}]
    )
    dim_question = pd.DataFrame(
        [{"question": q, "sort_order": i} for i, q in enumerate(TURN_ORDER)]
    )

    fact_summary = pd.DataFrame([summary_arm(df) for df in all_persona])
    fact_recall = pd.concat([recall_by_class(df) for df in all_persona], ignore_index=True)

    mcnemar_list = []
    pair_sheets: dict[str, pd.DataFrame] = {}
    comparisons = [
        ("openai_v1_vs_openai_v10", "openai_v1", "openai_v10", "mcnemar_oai_v1_v10"),
        ("claude_v1_vs_claude_v10", "claude_v1", "claude_v10", "mcnemar_claude_v1_v10"),
        ("openai_v1_vs_claude_v1", "openai_v1", "claude_v1", "mcnemar_oai_v1_claude_v1"),
        ("openai_v10_vs_claude_v10", "openai_v10", "claude_v10", "mcnemar_oai10_claude10"),
    ]
    for comp_id, a_id, b_id, sheet in comparisons:
        if a_id not in personas_by_arm or b_id not in personas_by_arm:
            continue
        pairs, stats = mcnemar_stats(personas_by_arm[a_id], personas_by_arm[b_id], comp_id)
        stats["arm_a_id"] = a_id
        stats["arm_b_id"] = b_id
        mcnemar_list.append(stats)
        pair_sheets[sheet] = pairs

    mcnemar_results = pd.DataFrame(mcnemar_list)

    readme = pd.DataFrame(
        [
            {"topic": "initiative", "detail": "New Horizons: reduce FN, reduce coordinator MR, completion rates"},
            {"topic": "cohort", "detail": COHORT},
            {"topic": "personas", "detail": "F001-F120 x 4 arms"},
            {"topic": "FN flags", "detail": "fn_missed_* and unsafe_approve on fact_persona"},
            {"topic": "coordinator", "detail": "unnecessary_manual_review on fact_persona"},
            {"topic": "completion", "detail": "screening_completed on fact_persona"},
            {"topic": "criteria", "detail": "fact_rationale scope=criterion_flag"},
            {"topic": "templates", "detail": "tpl_eval_screenings, tpl_actual_chatbot (empty)"},
            {"topic": "encoding", "detail": "0=ineligible 1=eligible 2=insufficient 3=manual_review"},
        ]
    )

    sheets: dict[str, pd.DataFrame] = {
        "README": readme,
        "tpl_eval_screenings": template_eval_screenings(),
        "tpl_actual_chatbot": template_actual_chatbot_screenings(),
        "dim_encoding": dim_encoding,
        "dim_arm": dim_arm,
        "dim_cohort": dim_cohort,
        "dim_question": dim_question,
        "fact_persona": fact_persona,
        "fact_turn": fact_turn,
        "fact_checkpoint": fact_checkpoint,
        "fact_rationale": fact_rationale,
        "fact_summary_arm": fact_summary,
        "fact_recall_by_class": fact_recall,
        "mcnemar_results": mcnemar_results,
        **pair_sheets,
    }

    for out_dir in OUT_DIRS:
        if not str(out_dir):
            continue
        xlsx = out_dir / XLSX_NAME
        docx = out_dir / DOCX_NAME
        write_xlsx(xlsx, sheets)
        write_docx(docx, mcnemar_results)
        print("Wrote", xlsx, "rows persona", len(fact_persona), "turn", len(fact_turn))
        print("Wrote", docx)


if __name__ == "__main__":
    main()
